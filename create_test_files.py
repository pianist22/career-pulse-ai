# create_test_files.py
# Generates three sample resumes under data/raw:
#  - sample_resume_docx.docx (Word file)
#  - sample_resume_pdf.pdf (PDF with inserted text)
#  - sample_resume_scanned.jpg (image with multiline text for OCR)

from pathlib import Path

def ensure_dirs():
    raw_dir = Path("data/raw")
    raw_dir.mkdir(parents=True, exist_ok=True)
    return raw_dir

def make_docx(raw_dir: Path):
    from docx import Document  # python-docx
    doc = Document()
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("Email: john.doe@example.com | Phone: +1-555-123-4567")
    doc.add_paragraph("LinkedIn: linkedin.com/in/johndoe")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Detail-oriented software engineer with 3+ years of experience in backend systems and APIs.")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Software Engineer, Acme Corp (2022–Present)")
    doc.add_paragraph("Built and maintained Python microservices; improved API latency by 25%.", style="List Bullet")
    doc.add_paragraph("Implemented CI/CD pipelines and unit tests (>90% coverage).", style="List Bullet")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.Tech in Computer Science, MAIT")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Django, Flask, REST, SQL, Docker, AWS")

    path = raw_dir / "sample_resume_docx.docx"
    doc.save(path.as_posix())
    return path

def make_pdf(raw_dir: Path):
    import fitz  # PyMuPDF
    doc_pdf = fitz.open()
    page = doc_pdf.new_page()  # A new blank page
    # Text box rectangle: (left, top, right, bottom) points ~ 72ppi default user space
    rect = fitz.Rect(72, 72, 540, 770)
    pdf_text = """John Doe
Email: john.doe@example.com | Phone: +1-555-123-4567
LinkedIn: linkedin.com/in/johndoe

Summary
Results-driven data analyst with 2+ years in SQL, Python, and dashboards.

Experience
Data Analyst, Example Ltd. (2023–Present)
- Built ETL jobs and optimized SQL queries (30% faster dashboards).
- Automated weekly KPI reporting with Python and Pandas.

Education
B.Tech in Computer Science, MAIT

Skills
Python, Pandas, NumPy, SQL, Tableau, Git
"""
    # Insert wrapped text in a rectangle (left-aligned)
    page.insert_textbox(rect, pdf_text, fontsize=11, fontname="helv", align=0)
    path = raw_dir / "sample_resume_pdf.pdf"
    doc_pdf.save(path.as_posix())
    return path

def make_scanned_image(raw_dir: Path):
    from PIL import Image, ImageDraw, ImageFont
    # Create a white canvas
    img = Image.new("RGB", (1200, 1600), "white")
    draw = ImageDraw.Draw(img)
    image_text = (
        "John Doe\n"
        "Email: john.doe@example.com | Phone: +1-555-123-4567\n"
        "LinkedIn: linkedin.com/in/johndoe\n\n"
        "Summary\n"
        "Aspiring ML engineer with internships in NLP and model training.\n\n"
        "Experience\n"
        "ML Intern, Startup X (2024)\n"
        "- Fine-tuned BERT for text classification; improved F1 by 6%.\n"
        "- Built data preprocessing pipelines (spaCy, NLTK).\n\n"
        "Education\n"
        "B.Tech in Computer Science, MAIT\n\n"
        "Skills\n"
        "Python, scikit-learn, Transformers, spaCy, TensorFlow, Docker\n"
    )
    # Use a default bitmap font to avoid OS font dependencies
    font = ImageFont.load_default()
    draw.multiline_text((50, 50), image_text, fill=(0, 0, 0), font=font, spacing=4, align="left")
    path = raw_dir / "sample_resume_scanned.jpg"
    img.save(path.as_posix(), quality=95)
    return path

if __name__ == "__main__":
    raw_dir = ensure_dirs()
    try:
        docx_p = make_docx(raw_dir)
    except Exception as e:
        docx_p = None
        print("DOCX generation failed. Install python-docx and retry. Error:", e)

    try:
        pdf_p = make_pdf(raw_dir)
    except Exception as e:
        pdf_p = None
        print("PDF generation failed. Ensure PyMuPDF (fitz) is installed. Error:", e)

    try:
        img_p = make_scanned_image(raw_dir)
    except Exception as e:
        img_p = None
        print("Image generation failed. Ensure Pillow is installed. Error:", e)

    print("Created test files:")
    if docx_p: print(" -", docx_p)
    if pdf_p:  print(" -", pdf_p)
    if img_p:  print(" -", img_p)
